from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
import pandas as pd
import json
import os
import sys
from docx.shared import Cm, Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer
import yaml


# --------------------------------------------------
# Utility
# --------------------------------------------------


def load_config(config_path: str = "configs/config_report.yaml") -> dict:
    with open(config_path, "r") as f:
        return yaml.safe_load(f)


def set_font(doc, font_name="Times New Roman", font_size=12):
    # Paragrafi
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    # Tabelle
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    # Stili
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(font_size)


def replace_placeholders(doc, replacements: dict):

    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for run in p.runs:
                for k, v in replacements.items():
                    if k in run.text:
                        run.text = run.text.replace(k, v)

    # Corpo documento
    replace_in_paragraphs(doc.paragraphs)

    # Tabelle nel corpo
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    # Header e footer
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)


# --------------------------------------------------
# CONFIG
# --------------------------------------------------

cliente = "A4"
template_path = "templates/A4_Template.docx"

opere = {
    "P001_Sommacampagna": {
        "label": "P001 - SOMMACAMPAGNA",
        "comune": "SOMMACAMPAGNA (VR)",
    },
    "P002_Giuliari_Milani": {
        "label": "P002 - GIULIARI MILANI",
        "comune": "VERONA (VR)",
    },
    "P003_Gua": {"label": "P003 - GUA", "comune": "GUA (VR)"},
    "P004_Adige_Est": {"label": "P004 - ADIGE EST", "comune": "VERONA (VR)"},
    "P005_Adige_Ovest": {"label": "P005 - ADIGE OVEST", "comune": "VERONA (VR)"},
}

OPERE_TO_KEY = {
    "P001": "P001_Sommacampagna",
    "P002": "P002_Giuliari_Milani",
    "P003": "P003_Gua",
    "P004": "P004_Adige_Est",
    "P005": "P005_Adige_Ovest",
}

MESI_IT = {
    1: "GENNAIO",
    2: "FEBBRAIO",
    3: "MARZO",
    4: "APRILE",
    5: "MAGGIO",
    6: "GIUGNO",
    7: "LUGLIO",
    8: "AGOSTO",
    9: "SETTEMBRE",
    10: "OTTOBRE",
    11: "NOVEMBRE",
    12: "DICEMBRE",
}


# --------------------------------------------------
# MAIN
# --------------------------------------------------


def run_report(
    config_path: str = "configs/config_report.yaml",
    year: str | None = None,
    month: str | None = None,
) -> None:
    # --------------------------------------------------
    # Resolve opera, year, month from config
    # --------------------------------------------------
    config = load_config(config_path)
    site_code = config["site"]["code"]
    ym = config["data"]["month"]

    if year is None:
        year = ym[:4]
    if month is None:
        month = ym[5:7]

    opera_key = OPERE_TO_KEY.get(site_code, f"{site_code}_Unknown")
    if opera_key not in opere:
        opera_key = f"{site_code}_Unknown"
        opera_info = {"label": site_code, "comune": ""}
    else:
        opera_info = opere[opera_key]

    opera_label = opera_info["label"]
    opera_comune = opera_info["comune"]

    year_int = int(year)
    month_int = int(month)
    mese_tag = f"{year}_{month}"
    mese_nome = f"{MESI_IT[month_int]} {year}"
    start_month = date(year_int, month_int, 1)
    end_month = start_month + relativedelta(months=1)

    # --------------------------------------------------
    # Paths and inputs
    # --------------------------------------------------
    fig_dir = os.path.join("figures", opera_key, mese_tag)
    base_out = os.path.join("outputs", opera_key, mese_tag)
    summary_csv = os.path.join(fig_dir, f"{year}_{month}_summary.csv")

    if not os.path.isdir(fig_dir):
        print(f"  ⚠ figure mancanti ({fig_dir}), salto {mese_tag}")
        return
    if not os.path.exists(summary_csv):
        print(f"  ⚠ summary mancante ({summary_csv}), salto {mese_tag}")
        return

    summary_df = pd.read_csv(summary_csv)

    # Optional KPIs and metrics (for future use)
    kpis_path = os.path.join(base_out, "kpis.json")
    if os.path.exists(kpis_path):
        with open(kpis_path, encoding="utf-8") as f:
            kpis = json.load(f)
    else:
        kpis = {
            "avg_displacement": "N/A",
            "max_acceleration": "N/A",
            "n_anomalies": int(summary_df["alarms"].sum()) if "alarms" in summary_df.columns else 0,
        }

    metrics_path = os.path.join(base_out, "model_metrics.csv")
    metrics_df = pd.read_csv(metrics_path) if os.path.exists(metrics_path) else None

    all_figures = sorted(
        f for f in os.listdir(fig_dir) if f.lower().endswith((".png", ".jpg", ".jpeg"))
    )
    figures_no_torte = [f for f in all_figures if "torte" not in f.lower()]

    # --------------------------------------------------
    # Create document from template + description
    # --------------------------------------------------
    replacements = {
        "Cliente": cliente,
        "{{data}}": datetime.today().strftime("%d/%m/%Y"),
        "{{MESE_ANNO}}": mese_nome,
        "{{PERIODO_DAL}}": start_month.strftime("%d/%m/%Y"),
        "{{PERIODO_AL}}": end_month.strftime("%d/%m/%Y"),
        "{{OPERA}}": opera_label,
        "{{Comune}}": opera_comune,
    }

    master = Document(template_path) if os.path.exists(template_path) else Document()
    replace_placeholders(master, replacements)

    # Append description doc (if available) using Composer
    composer = Composer(master)
    desc_prefix = opera_label[:4]  # e.g. "P005"
    desc_path = os.path.join("templates", "description", f"{desc_prefix}_description.docx")
    if os.path.exists(desc_path):
        description = Document(desc_path)
        composer.append(description)
        composer.save("temp.docx")
    doc = master  # continue building into master

    doc.add_page_break()

    # --------------------------------------------------
    # Section 2: Disponibilità Dati (torte plot)
    # --------------------------------------------------
    doc.add_heading("2. Disponibilità Dati", level=1)

    doc.add_paragraph("\n")
    
    df = pd.read_csv("outputs/nans_percentage.csv")
    df = df[['sensore', 'label', 'dati mancanti']]
    df = df.rename(columns={'sensore': 'ID sensore', 'label': 'Label', 'dati mancanti': 'Dati mancanti'})

    n_rows, n_cols = df.shape
    col_widths = [Cm(2)] * n_cols

    # rows first, then columns
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.autofit = False

    # set column widths
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # header row
    for j, col_name in enumerate(df.columns):
        table.cell(0, j).text = str(col_name)

    # table body
    for i in range(n_rows):
        for j in range(n_cols):
            value = df.iat[i, j]

            if df.columns[j] == "Dati mancanti":
                text = f"{float(value):.2f}%"
            else:
                text = str(value)

            table.cell(i + 1, j).text = text


    # for i, width in enumerate(col_widths):
    #     table.columns[i].width = width
    #     for cell in table.columns[i].cells:
    #         tc = cell._tc
    #         tcPr = tc.get_or_add_tcPr()
    #         w = OxmlElement("w:tcW")
    #         w.set(qn("w:w"), str(int(width.pt * 30)))
    #         w.set(qn("w:type"), "dxa")
    #         tcPr.append(w)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # for i, col in enumerate(summary_df.columns):
    #     table.rows[0].cells[i].text = str(col)

    # for _, row in summary_df.iterrows():
    #     row_cells = table.add_row().cells
    #     for i, val in enumerate(row):
    #         row_cells[i].text = str(val)


    '''torte_file = None
                for f in all_figures:
                    if "torte" in f.lower() and f.lower().endswith(".png"):
                        torte_file = os.path.join(fig_dir, f)
                        break
            
                if torte_file:
                    doc.add_paragraph()
                    doc.add_picture(torte_file, width=Inches(6.5))
                else:
                    doc.add_paragraph("Torte plot not available for this period.")'''

    doc.add_page_break()

    # --------------------------------------------------
    # Section 3: Visualizzazione Dati Grezzi
    # --------------------------------------------------
    doc.add_heading("3. Visualizzazione Dati Grezzi", level=1)
    doc.add_paragraph(
        "Si riportano di seguito i grafici dei dati grezzi suddivisi per tipologia "
        "di sensore, relativi al periodo di riferimento di questo report.\n"
        "All'andamento di ogni sensore è affiancato l'andamento della temperatura "
        "di una sonda di riferimento, per evidenziare eventuali correlazioni.\n"
    )

    tipologie_raw = {
        "Inclinometri": "raw_ICD",
        "Potenziometri": "raw_POT",
        "Estensimetri": "raw_EST",
    }

    for nome_tipologia, prefisso in tipologie_raw.items():
        tipo_figures = sorted(f for f in figures_no_torte if f.startswith(prefisso))
        if not tipo_figures:
            continue

        doc.add_heading(nome_tipologia, level=2)
        table = doc.add_table(rows=1, cols=2)

        for i in range(0, len(tipo_figures), 2):
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].add_run().add_picture(
                os.path.join(fig_dir, tipo_figures[i]), width=Cm(9)
            )
            if i < len(tipo_figures) - 1:
                row_cells[1].paragraphs[0].add_run().add_picture(
                    os.path.join(fig_dir, tipo_figures[i + 1]), width=Cm(9)
                )
        doc.add_paragraph()

    if not any(f.startswith("raw_") for f in figures_no_torte):
        doc.add_paragraph("Non sono disponibili grafici dei sensori per il mese selezionato.")

    doc.add_page_break()

    # --------------------------------------------------
    # Section 4: Visualizzazione Z-Score
    # --------------------------------------------------
    doc.add_heading("4. Visualizzazione Z-Score", level=1)
    doc.add_paragraph(
        "Si riportano di seguito i grafici degli z-score suddivisi per tipologia "
        "di sensore, relativi al periodo di riferimento di questo report.\n"
        "Per z-score si intende il numero di deviazioni standard con cui "
        "una misura si discosta dalla relativa media.\n"
    )

    def classifica_tipologia(nome_file: str) -> str | None:
        base = os.path.splitext(nome_file)[0]
        suffisso = base[-1].lower()
        if not 'z-score' in nome_file:
            return None
        elif suffisso in ["x", "y"]:
            return "Inclinometri"
        elif suffisso == "s":
            return "Potenziometri"
        elif suffisso == "e":
            return "Estensimetri"
        else:
            return None

    tipologie_zscore: dict[str, list[str]] = {
        "Inclinometri": [],
        "Potenziometri": [],
        "Estensimetri": [],
    }
    for f in figures_no_torte:
        categoria = classifica_tipologia(f)
        if categoria:
            tipologie_zscore[categoria].append(f)

    for nome_tipologia, tipo_figures in tipologie_zscore.items():
        tipo_figures = sorted(tipo_figures)
        if not tipo_figures:
            continue

        doc.add_heading(nome_tipologia, level=2)
        table = doc.add_table(rows=1, cols=2)
        for i in range(0, len(tipo_figures), 2):
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].add_run().add_picture(
                os.path.join(fig_dir, tipo_figures[i]), width=Cm(9)
            )
            if i < len(tipo_figures) - 1:
                row_cells[1].paragraphs[0].add_run().add_picture(
                    os.path.join(fig_dir, tipo_figures[i + 1]), width=Cm(9)
                )
        doc.add_paragraph()

    doc.add_page_break()

    # --------------------------------------------------
    # Section 5: Warnings e Alerts (summary table)
    # --------------------------------------------------

    summary_df = summary_df.rename(columns={'sensor_id': 'ID sensore', 'label': 'Label', 'warnings': 'Warnings', 'alarms': 'Allarmi' })
    doc.add_heading("5. Warnings e Allarmi", level=1)

    doc.add_paragraph(
            "\nSi riporta di seguito una tabella contenente, per ciascun sensore, il numero di superamenti delle soglie di controllo (warnings) e di allarme."
            "Le soglie di controllo sono state fissate ad un valore pari a tre deviazioni standard attorno al valore medio del segnale corrispondente."
            "La soglia di allarme è stata definita come da documento inviato in data 28.11.2025: 'Comunicazione gestione soglie di allarme' \n"
    )

    n_cols = len(summary_df.columns)
    col_widths = [Cm(2)] * n_cols

    table = doc.add_table(rows=1, cols=n_cols)
    table.autofit = False

    for i, width in enumerate(col_widths):
        table.columns[i].width = width
        for cell in table.columns[i].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            w = OxmlElement("w:tcW")
            w.set(qn("w:w"), str(int(width.pt * 30)))
            w.set(qn("w:type"), "dxa")
            tcPr.append(w)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(summary_df.columns):
        table.rows[0].cells[i].text = str(col)

    for _, row in summary_df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.add_page_break()

    # --------------------------------------------------
    # Save
    # --------------------------------------------------
    os.makedirs(base_out, exist_ok=True)
    output_name = f"{cliente}_{opera_key}_{mese_tag}.docx"
    output_path = os.path.join(base_out, output_name)
    doc.save(output_path)
    print(f"\033[92m✔ salvato {output_name}\033[0m")


if __name__ == "__main__":
    if len(sys.argv) >= 3:
        run_report(year=sys.argv[1], month=sys.argv[2])
    elif len(sys.argv) == 2:
        run_report(config_path=sys.argv[1])
    else:
        run_report()
