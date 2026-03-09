"""Report generation step: DOCX from templates and figures."""

import json
import os
from datetime import date, datetime

import pandas as pd
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docxcompose.composer import Composer

from timeseries_etl.config import load_config
from timeseries_etl.domain import MESI_IT, get_opera_info
from timeseries_etl.domain.labels import load_label_dict

CLIENTE = "A4"
TEMPLATE_PATH = "templates/A4_Template.docx"
NANS_CSV = "outputs/nans_percentage.csv"


def _replace_placeholders(doc: Document, replacements: dict[str, str]) -> None:
    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for run in p.runs:
                for k, v in replacements.items():
                    if k in run.text:
                        run.text = run.text.replace(k, str(v))

    replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)


def _classifica_tipologia_zscore(nome_file: str) -> str | None:
    base = os.path.splitext(nome_file)[0]
    suffisso = base[-1].lower() if base else ""
    if "z-score" not in nome_file.lower():
        return None
    if suffisso in ["x", "y"]:
        return "Inclinometri"
    if suffisso == "s":
        return "Potenziometri"
    if suffisso == "e":
        return "Estensimetri"
    return None


def run_report(
    config_path: str = "configs/config_report.yaml",
    year: str | None = None,
    month: str | None = None,
) -> None:
    """Generate DOCX report from templates, figures, and CSVs."""
    config = load_config(config_path)
    site_code = config["site"]["code"]
    ym = config["data"]["month"]

    year = year or ym[:4]
    month = month or ym[5:7]

    opera_key, opera_info = get_opera_info(site_code)

    opera_label = opera_info["label"]
    opera_comune = opera_info["comune"]

    year_int = int(year)
    month_int = int(month)
    mese_tag = f"{year}_{month}"
    mese_nome = f"{MESI_IT[month_int]} {year}"
    start_month = date(year_int, month_int, 1)
    end_month = start_month + relativedelta(months=1)

    fig_dir = os.path.join("figures", site_code, mese_tag)
    base_out = os.path.join("outputs", site_code, mese_tag)
    summary_csv = os.path.join(fig_dir, f"{year}_{month}_summary.csv")

    if not os.path.isdir(fig_dir):
        print(f"  ⚠ figure mancanti ({fig_dir}), salto {mese_tag}")
        return
    if not os.path.exists(summary_csv):
        print(f"  ⚠ summary mancante ({summary_csv}), salto {mese_tag}")
        return

    summary_df = pd.read_csv(summary_csv)

    kpis_path = os.path.join(base_out, "kpis.json")
    if os.path.exists(kpis_path):
        with open(kpis_path, encoding="utf-8") as f:
            kpis = json.load(f)
    else:
        n_anomalies = int(summary_df["alarms"].sum()) if "alarms" in summary_df.columns else 0
        kpis = {"avg_displacement": "N/A", "max_acceleration": "N/A", "n_anomalies": n_anomalies}

    _ = kpis  # reserved for future use
    all_figures = sorted(
        f for f in os.listdir(fig_dir) if f.lower().endswith((".png", ".jpg", ".jpeg"))
    )
    figures_no_torte = [f for f in all_figures if "torte" not in f.lower()]

    replacements = {
        "Cliente": CLIENTE,
        "{{data}}": datetime.today().strftime("%d/%m/%Y"),
        "{{MESE_ANNO}}": mese_nome,
        "{{PERIODO_DAL}}": start_month.strftime("%d/%m/%Y"),
        "{{PERIODO_AL}}": end_month.strftime("%d/%m/%Y"),
        "{{OPERA}}": opera_label,
        "{{Comune}}": opera_comune,
    }

    master = Document(TEMPLATE_PATH) if os.path.exists(TEMPLATE_PATH) else Document()
    _replace_placeholders(master, replacements)

    composer = Composer(master)
    desc_prefix = opera_label[:4]
    desc_path = os.path.join("templates", "description", f"{desc_prefix}_description.docx")
    if os.path.exists(desc_path):
        description = Document(desc_path)
        composer.append(description)
        composer.save("temp.docx")

    doc = master
    doc.add_page_break()

    doc.add_heading("2. Disponibilità Dati", level=1)
    doc.add_paragraph("\n")

    if os.path.exists(NANS_CSV):
        nans_df = pd.read_csv(NANS_CSV)
        nans_df = nans_df[["sensore", "label", "dati mancanti"]]
        nans_df = nans_df.sort_values(by="label")
        nans_df = nans_df.rename(
            columns={"sensore": "ID sensore", "label": "Label", "dati mancanti": "Dati mancanti"}
        )
        n_rows, n_cols = nans_df.shape
        col_widths = [Cm(2)] * n_cols
        table = doc.add_table(rows=n_rows + 1, cols=n_cols)
        table.autofit = False
        for i, width in enumerate(col_widths):
            table.columns[i].width = width
        for j, col_name in enumerate(nans_df.columns):
            table.cell(0, j).text = str(col_name)
        for i in range(n_rows):
            for j in range(n_cols):
                value = nans_df.iat[i, j]
                text = f"{float(value):.2f}%" if nans_df.columns[j] == "Dati mancanti" else str(value)
                table.cell(i + 1, j).text = text
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

    doc.add_page_break()

    doc.add_heading("3. Visualizzazione Dati Grezzi", level=1)
    doc.add_paragraph(
        "Si riportano di seguito i grafici dei dati grezzi suddivisi per tipologia "
        "di sensore, relativi al periodo di riferimento di questo report.\n"
        "All'andamento di ogni sensore è affiancato l'andamento della temperatura "
        "di una sonda di riferimento, per evidenziare eventuali correlazioni.\n"
    )

    label_dict = load_label_dict(site_code)

    def _raw_fig_sort_key(f: str) -> tuple[str, str]:
        """Sort raw figures by (label, axis_tag)."""
        base = f.replace("raw_", "").replace(".png", "").replace(".jpg", "").replace(".jpeg", "")
        for suffix in ["_x", "_y", "_s", "_e"]:
            if base.endswith(suffix):
                return (base[:-len(suffix)], suffix)
        return (base, "")

    tipologie_raw = {
        "Inclinometri": "raw_ICD",
        "Potenziometri": "raw_POT",
        "Estensimetri": "raw_EST",
    }
    for nome_tipologia, prefisso in tipologie_raw.items():
        tipo_figures = sorted(
            (f for f in figures_no_torte if f.startswith(prefisso)),
            key=_raw_fig_sort_key,
        )
        if not tipo_figures:
            continue
        doc.add_heading(nome_tipologia, level=2)
        table = doc.add_table(rows=1, cols=2)
        for i in range(0, len(tipo_figures), 2):
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].add_run().add_picture(
                os.path.join(fig_dir, tipo_figures[i]), width=Cm(9)
            )
            if i < len(tipo_figures) - 1:
                row_cells[1].paragraphs[0].add_run().add_picture(
                    os.path.join(fig_dir, tipo_figures[i + 1]), width=Cm(9)
                )
        doc.add_paragraph()

    if not any(f.startswith("raw_") for f in figures_no_torte):
        doc.add_paragraph("Non sono disponibili grafici dei sensori per il mese selezionato.")

    doc.add_page_break()

    doc.add_heading("4. Visualizzazione Z-Score", level=1)
    doc.add_paragraph(
        "Si riportano di seguito i grafici degli z-score suddivisi per tipologia "
        "di sensore, relativi al periodo di riferimento di questo report.\n"
        "Per z-score si intende il numero di deviazioni standard con cui "
        "una misura si discosta dalla relativa media.\n"
    )

    tipologie_zscore: dict[str, list[str]] = {
        "Inclinometri": [],
        "Potenziometri": [],
        "Estensimetri": [],
    }
    for f in figures_no_torte:
        cat = _classifica_tipologia_zscore(f)
        if cat:
            tipologie_zscore[cat].append(f)

    def _zscore_fig_sort_key(f: str) -> str:
        """Sort z-score figures by label (from summary_df)."""
        sensor_id = f.replace("z-score_", "").rsplit(".", 1)[0]
        row = summary_df[summary_df["sensor_id"] == sensor_id]
        return row["label"].iloc[0] if len(row) else sensor_id

    for nome_tipologia, tipo_figures in tipologie_zscore.items():
        tipo_figures = sorted(tipo_figures, key=_zscore_fig_sort_key)
        if not tipo_figures:
            continue
        doc.add_heading(nome_tipologia, level=2)
        table = doc.add_table(rows=1, cols=2)
        for i in range(0, len(tipo_figures), 2):
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].add_run().add_picture(
                os.path.join(fig_dir, tipo_figures[i]), width=Cm(9)
            )
            if i < len(tipo_figures) - 1:
                row_cells[1].paragraphs[0].add_run().add_picture(
                    os.path.join(fig_dir, tipo_figures[i + 1]), width=Cm(9)
                )
        doc.add_paragraph()

    doc.add_page_break()

    summary_df = summary_df.sort_values(by="label")
    summary_df = summary_df.rename(
        columns={
            "sensor_id": "ID sensore",
            "label": "Label",
            "warnings": "Warnings",
            "alarms": "Allarmi",
        }
    )
    doc.add_heading("5. Warnings e Allarmi", level=1)
    doc.add_paragraph(
        "\nSi riporta di seguito una tabella contenente, per ciascun sensore, il numero di "
        "superamenti delle soglie di controllo (warnings) e di allarme. "
        "Le soglie di controllo sono state fissate ad un valore pari a tre deviazioni standard "
        "attorno al valore medio del segnale corrispondente. "
        "La soglia di allarme è stata definita come da documento inviato in data 28.11.2025: "
        "'Comunicazione gestione soglie di allarme' \n"
    )

    n_cols = len(summary_df.columns)
    col_widths = [Cm(2)] * n_cols
    table = doc.add_table(rows=1, cols=n_cols)
    table.autofit = False
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
        for cell in table.columns[i].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            w = OxmlElement("w:tcW")
            w.set(qn("w:w"), str(int(width.pt * 30)))
            w.set(qn("w:type"), "dxa")
            tcPr.append(w)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(summary_df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in summary_df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.add_page_break()

    os.makedirs(base_out, exist_ok=True)
    output_name = f"{CLIENTE}_{site_code}_{mese_tag}.docx"
    output_path = os.path.join(base_out, output_name)
    doc.save(output_path)
    print(f"\033[92m✔ salvato {output_name}\033[0m")
